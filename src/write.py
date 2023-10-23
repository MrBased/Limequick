from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


template_path = "TemplateLime.docx"
new_file_path = "essay_code.docx"

document = Document(template_path)
document.save(new_file_path)


new_document = Document(new_file_path)


insert_paragraph_index1 = 8
insert_paragraph_index2 = 13


def insert_text_with_bullet(document, paragraph_index, text):
    paragraph = document.paragraphs[paragraph_index]

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run = paragraph.add_run("- ")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.add_text(" ")
    run.add_text(text)
    paragraph.add_run("\n")


n = 0
n_good_comments = int(input("How many positive comments do you want to add? "))
while n < n_good_comments:
    insert_text_with_bullet(new_document, insert_paragraph_index1, input("Your comment:"))
    n += 1


n = 0
n_bad_comments = int(input("How many suggestions for improvement do you want to add? "))
while n < n_bad_comments:
    insert_text_with_bullet(new_document, insert_paragraph_index2, input("Your comment:"))
    n += 1

new_document.save(new_file_path)

print("File saved!")
