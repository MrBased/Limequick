from flask import Flask, render_template, request, send_file
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import os


insert_paragraph_index1 = 8
insert_paragraph_index2 = 13

app = Flask(__name__)


app.config['UPLOAD_FOLDER'] = 'uploads'
uploads_dir = os.path.join(app.root_path, 'uploads')

if not os.path.exists(uploads_dir):
    os.makedirs(uploads_dir)


def insert_text_with_bullet(document, paragraph_index, text):
    paragraph = document.paragraphs[paragraph_index]

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run = paragraph.add_run("- ")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.add_text(" ")
    run.add_text(text)
    paragraph.add_run("\n")


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        n_good_comments = int(request.form['n_good_comments'])
        n_bad_comments = int(request.form['n_bad_comments'])

        new_file_path = "src/static/result.docx"
        template_path = "src/TemplateLime.docx"

        new_document = Document(template_path)

        for _ in range(n_good_comments):
            comment = request.form.get('good_comment')
            insert_text_with_bullet(new_document, insert_paragraph_index1, comment)

        for _ in range(n_bad_comments):
            comment = request.form.get('bad_comment')
            insert_text_with_bullet(new_document, insert_paragraph_index2, comment)

        new_document.save(new_file_path)
        # pdf_path = "static/essay_code.pdf"

        return render_template('index.html', pdf_created=True)
    return render_template('index.html', pdf_created=False)


@app.route('/download_pdf')
def download_pdf():
    pdf_path = "static/essay_code.docx"
    return send_file(pdf_path, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
